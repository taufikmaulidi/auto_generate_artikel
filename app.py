import os
import google.generativeai as genai
import requests
import streamlit as st
import docx
import zipfile
from io import BytesIO
import re

# Validasi API Key Google Gemini
def validate_google_gemini_api_key(api_key):
    headers = {
        'Content-Type': 'application/json',
    }

    params = {
        'key': api_key,
    }

    json_data = {
        'contents': [
            {
                'role': 'user',
                'parts': [
                    {
                        'text': 'Give me five subcategories of jazz?',
                    },
                ],
            },
        ],
    }

    response = requests.post(
        'https://generativelanguage.googleapis.com/v1/models/gemini-pro:generateContent',
        params=params,
        headers=headers,
        json=json_data,
    )
    return response.status_code == 200

# Fungsi untuk membersihkan teks dari simbol markdown dengan sensitivitas tinggi
def clean_text(text):
    # Mengganti simbol heading '#' dengan format yang sesuai
    text = re.sub(r'^#\s*(.*)', r'\1', text, flags=re.MULTILINE)  # Mengganti heading level 1
    text = re.sub(r'^\s*##\s*(.*)', r'\1', text, flags=re.MULTILINE)  # Mengganti heading level 2
    text = re.sub(r'^\s*###\s*(.*)', r'\1', text, flags=re.MULTILINE)  # Mengganti heading level 3
    
    # Mengganti bullet points '*' dan '-' dengan dash '-' yang konsisten
    text = re.sub(r'^\s*\*\s*(.*)', r'- \1', text, flags=re.MULTILINE)  # Bullet points dengan '*'
    text = re.sub(r'^\s*-\s*(.*)', r'- \1', text, flags=re.MULTILINE)  # Bullet points dengan '-'
    
    # Menghapus simbol '*' di tempat lain, kecuali yang digunakan untuk bold
    text = re.sub(r'\*{2}(.*?)\*{2}', r'\1', text)  # Mengganti bold text dengan normal text
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Mengganti italic text dengan normal text

    # Menghapus simbol '#' dan '*' yang mungkin tersisa di bagian lain
    text = re.sub(r'[\#\*]', '', text)
    
    # Mengganti dua baris kosong dengan satu baris kosong
    text = re.sub(r'\n\s*\n', '\n\n', text)

    return text

# Save response to a DOCX file
def save_to_docx(title, content):
    doc = docx.Document()
    doc.add_heading(title, level=1)
    clean_content = clean_text(content)
    doc.add_paragraph(clean_content)
    file_name = f"{title.replace(' ', '_')}.docx"
    doc.save(file_name)
    return file_name

# Set page config
st.set_page_config(page_title='Auto Generate Artikel',
                   page_icon='./favicon.png', 
                   layout='wide',
                   )

# Custom CSS to hide Streamlit footer and GitHub icon
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            .css-1rs6os {visibility: hidden;}  /* Hide GitHub icon */
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# Initialize session state for API key
if 'api_key_valid' not in st.session_state:
    st.session_state.api_key_valid = False
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""
st.title('Generate Artikel')
st.subheader("Buat Artikel Mudah dengan AI")
with st.sidebar:
    st.title("Auto Generate Artikel")
    st.subheader('by gudanginformatika.com')

    if not st.session_state.api_key_valid:
        api_key = st.text_input('Masukkan API Key Google Gemini Anda', type='password')
        if st.button('Validate API Key'):
            if validate_google_gemini_api_key(api_key):
                st.session_state.api_key_valid = True
                st.session_state.api_key = api_key
                st.success('API Key valid!')
            else:
                st.error('API Key tidak valid')

if st.session_state.api_key_valid:
    genai.configure(api_key=st.session_state.api_key)
    generation_config = {
        "temperature": 1,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }

    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        generation_config=generation_config,
    )

    # Tampilkan Form Generate Artikel
    gaya_bahasa = st.radio("Gaya Bahasa",
                           ["Informatif", "Naratif", "Kasual","Formal","Kreatif"],
                           captions=["Memberikan informasi yang akurat dan bermanfaat kepada pembaca",
                                     "Menceritakan sebuah kisah yang menarik dan engaging bagi pembaca.",
                                     "Meyakinkan pembaca untuk mengambil tindakan tertentu, seperti membeli produk, mendaftar newsletter, atau mendukung suatu opini.",
                                     "Menciptakan suasana yang santai dan bersahabat dengan pembaca.",
                                     "Menyampaikan informasi yang serius dan kredibel kepada pembaca.",
                                     "Menyampaikan informasi dengan cara yang unik dan imajinatif."])
    num_len = st.slider("Length of Words", min_value=500, max_value=2000, step=100)

    uploaded_file = st.file_uploader("Upload TXT File", type="txt")
    if uploaded_file is not None:
        text_data = uploaded_file.read().decode("utf-8")
        titles = text_data.splitlines()

        if st.button('Generate Articles'):
            docx_files = []
            for title in titles:
                input_prompt = f"""
                Anda adalah seorang SEO spesialis dengan pengalaman dalam membuat artikel SEO yang mudah diindeks mesin pencari dan menarik bagi pembaca. Tugas Anda adalah membuat artikel blog dengan judul artikel yang harus Anda buat adalah {title}. Jika judul ini tidak berstandar SEO, modifikasi judulnya sesuai teknik SEO untuk memastikan optimalisasi mesin pencari. Selanjutnya, tulis artikel blog dengan judul yang telah diberikan. Gunakan gaya penulisan {gaya_bahasa} untuk memastikan daya tarik yang tinggi. Artikel harus memiliki jumlah kata {num_len}. Pastikan artikel relevan, informatif, dan sesuai dengan standar SEO untuk memaksimalkan visibilitas di mesin pencari.
                """

                response = model.generate_content(input_prompt)
                if response:
                    content = response.text
                    file_name = save_to_docx(title, content)
                    docx_files.append(file_name)
                    st.success(f"Artikel '{title}' berhasil dibuat!")
            
            if len(docx_files) > 1:
                # Create a zip file
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                    for file_name in docx_files:
                        zip_file.write(file_name)
                zip_buffer.seek(0)
                
                st.download_button(
                    label="Download All Articles as ZIP",
                    data=zip_buffer.getvalue(),
                    file_name="articles.zip",
                    mime="application/zip"
                )
            elif len(docx_files) == 1:
                single_file = docx_files[0]
                st.download_button(
                    label="Download Article",
                    data=open(single_file, "rb").read(),
                    file_name=single_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Gagal membuat artikel.")
